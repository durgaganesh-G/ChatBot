# -------------------------
# Hide warnings
# -------------------------
import warnings
warnings.filterwarnings("ignore")

# -------------------------
# FastAPI
# -------------------------
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Dict

# -------------------------
# LangChain
# -------------------------
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.schema import Document

# -------------------------
# Transformers
# -------------------------
from transformers import AutoTokenizer, AutoModelForQuestionAnswering
import torch

import os

app = FastAPI()

# -------------------------
# Allow frontend requests
# -------------------------
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# -------------------------
# Global state
# -------------------------
db = None
conversation_history: List[Dict[str, str]] = []   # [{question, answer}, ...]

# -------------------------
# Request model
# -------------------------
class Question(BaseModel):
    question: str


# -------------------------
# Load models once at startup
# -------------------------
model_name = "distilbert-base-cased-distilled-squad"

tokenizer = AutoTokenizer.from_pretrained(model_name)
model = AutoModelForQuestionAnswering.from_pretrained(model_name)

embeddings = HuggingFaceEmbeddings(
    model_name="sentence-transformers/all-MiniLM-L6-v2"
)


# -------------------------
# Helper: extract text from uploaded file bytes
# -------------------------
def extract_text(filename: str, file_bytes: bytes) -> str:
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".txt":
        return file_bytes.decode("utf-8", errors="ignore")

    elif ext == ".pdf":
        try:
            import io
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            return text
        except ImportError:
            raise HTTPException(
                status_code=500,
                detail="PyPDF2 not installed. Run: pip install PyPDF2"
            )

    elif ext in (".docx", ".doc"):
        try:
            import io
            import docx
            doc = docx.Document(io.BytesIO(file_bytes))
            text = "\n".join([para.text for para in doc.paragraphs])
            return text
        except ImportError:
            raise HTTPException(
                status_code=500,
                detail="python-docx not installed. Run: pip install python-docx"
            )

    else:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type '{ext}'. Please upload a .txt, .pdf, or .docx file."
        )


# -------------------------
# Upload file (TXT, PDF, DOCX)
# -------------------------
@app.post("/upload")
async def upload_file(file: UploadFile = File(...)):

    global db, conversation_history

    # Read file bytes properly (async-safe)
    file_bytes = await file.read()

    if not file_bytes:
        raise HTTPException(status_code=400, detail="Uploaded file is empty.")

    # Extract text based on file type
    text = extract_text(file.filename, file_bytes)

    if not text.strip():
        raise HTTPException(
            status_code=400,
            detail="Could not extract any text from the uploaded file. Make sure it contains readable text."
        )

    # Wrap in a LangChain Document
    documents = [Document(page_content=text, metadata={"source": file.filename})]

    # Split text into chunks
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=50
    )
    texts = splitter.split_documents(documents)

    # Build vector DB from chunks
    db = FAISS.from_documents(texts, embeddings)

    # Clear conversation history for fresh session
    conversation_history = []

    return {
        "message": f"File '{file.filename}' uploaded and processed successfully. You can now ask questions."
    }


# -------------------------
# Ask a question (supports multiple questions in sequence)
# -------------------------
@app.post("/ask")
async def ask_question(data: Question):

    global db, conversation_history

    if db is None:
        raise HTTPException(status_code=400, detail="Upload a document first")

    query = data.question.strip()
    if not query:
        raise HTTPException(status_code=400, detail="Question cannot be empty")

    # Retrieve relevant chunks from vector DB
    docs = db.similarity_search(query, k=3)

    if not docs:
        return {"answer": "No relevant information found.", "history": conversation_history}

    # Build context: retrieved document chunks
    doc_context = " ".join([doc.page_content for doc in docs])

    # Add recent conversation history as additional context (last 3 turns)
    history_context = ""
    if conversation_history:
        recent = conversation_history[-3:]
        history_context = " ".join(
            [f"Q: {h['question']} A: {h['answer']}" for h in recent]
        )
        combined_context = history_context + " " + doc_context
    else:
        combined_context = doc_context

    # Tokenize and run model
    inputs = tokenizer(
        query,
        combined_context,
        return_tensors="pt",
        truncation=True,
        max_length=512
    )

    with torch.no_grad():
        outputs = model(**inputs)

    start = torch.argmax(outputs.start_logits)
    end = torch.argmax(outputs.end_logits) + 1

    answer = tokenizer.decode(inputs["input_ids"][0][start:end]).strip()

    if not answer or answer in ("[CLS]", "[SEP]", ""):
        answer = "Answer not found in the document."

    # Save this Q&A to conversation history
    conversation_history.append({
        "question": query,
        "answer": answer
    })

    return {
        "answer": answer,
        "question_number": len(conversation_history),
        "history": conversation_history
    }


# -------------------------
# Get full conversation history
# -------------------------
@app.get("/history")
async def get_history():
    return {
        "total_questions": len(conversation_history),
        "history": conversation_history
    }


# -------------------------
# Reset conversation history (without re-uploading file)
# -------------------------
@app.post("/reset")
async def reset_history():
    global conversation_history
    conversation_history = []
    return {"message": "Conversation history cleared. You can ask fresh questions on the same document."}